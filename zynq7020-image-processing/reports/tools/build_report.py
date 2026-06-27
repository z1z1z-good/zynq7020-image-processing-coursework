# -*- coding: utf-8 -*-
"""Build the ZYNQ7020 image-processing final coursework report as a .docx.

Source content: reports/最终报告.md (the user's final report).
Aligned with: teacher upstream README report requirements (section 7 最终报告) and
the latest 等级评定 grading tables (Task 1 -> C grade, Task 3 -> B grade).
New on-site acceptance images (2026-06-26) are integrated for the two extensions.
"""
import os
import re
from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

# ----------------------------------------------------------------------------
# paths — resolved relative to THIS script so the report can be regenerated
# from a fresh clone on any machine. Script lives at <repo>/.../reports/tools/.
# ----------------------------------------------------------------------------
HERE = os.path.dirname(os.path.abspath(__file__))
REPORTS = os.path.dirname(HERE)
ZYNQ = os.path.dirname(REPORTS)
EV = ZYNQ + "/coursework/evidence"
NEW = EV + "/照片证明/大拓展最终验收视频截图"
ASSETS = HERE + "/assets"
GDUT_LOGO = ASSETS + "/gdut_logo.png"
IC_LOGO = ASSETS + "/ic_school.png"
OUT = REPORTS + "/最终报告.docx"

CENTER = WD_ALIGN_PARAGRAPH.CENTER

# theme colors
NAVY = RGBColor(0x1F, 0x38, 0x64)
DARK = RGBColor(0x22, 0x22, 0x22)
GRAY = RGBColor(0x55, 0x55, 0x55)
HEADER_FILL = "1F3864"
SUBHEAD_FILL = "D9E2F3"
CODE_FILL = "F4F4F4"
CALL_FILL = "E7F0FF"

FIG = [0]
TBL = [0]

# ----------------------------------------------------------------------------
# image preprocessing: downscale oversized photos so the .docx stays small
# (display width is only 7-14 cm; 1600 px long edge is >250 DPI at that size).
# Tiny diagram/waveform PNGs are kept byte-for-byte to stay crisp.
# ----------------------------------------------------------------------------
TMP = "D:/Claude_prj/fpga/_imgtmp"
os.makedirs(TMP, exist_ok=True)
_prep_cache = {}
_prep_n = [0]


def prep(path):
    if path in _prep_cache:
        return _prep_cache[path]
    if not os.path.exists(path):
        return path
    try:
        size = os.path.getsize(path)
        im = Image.open(path)
        w, h = im.size
        long_edge = max(w, h)
        need_resize = long_edge > 1600
        if size < 250 * 1024 and not need_resize:
            _prep_cache[path] = path
            return path
        if im.mode != 'RGB':
            im = im.convert('RGB')
        if need_resize:
            scale = 1600.0 / long_edge
            im = im.resize((max(1, int(w * scale)), max(1, int(h * scale))), Image.LANCZOS)
        _prep_n[0] += 1
        outp = os.path.join(TMP, "img_%02d.jpg" % _prep_n[0])
        im.save(outp, "JPEG", quality=88, optimize=True)
        _prep_cache[path] = outp
        return outp
    except Exception as e:
        print("prep error", path, e)
        return path

# ----------------------------------------------------------------------------
# low level helpers
# ----------------------------------------------------------------------------

def set_font(run, latin, ea, size, bold=False, italic=False, color=None):
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:ascii'), latin)
    rfonts.set(qn('w:hAnsi'), latin)
    rfonts.set(qn('w:eastAsia'), ea)
    rfonts.set(qn('w:cs'), latin)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def shade_cell(cell, hex_):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_)
    tcPr.append(shd)


def shade_par(p, hex_):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_)
    pPr.append(shd)


def add_rich(p, text, size=10.5, base_color=DARK):
    """Render a small subset of markdown inline: `code` and **bold**."""
    tokens = re.split(r'(`[^`]*`|\*\*[^*]+\*\*)', text)
    for tok in tokens:
        if not tok:
            continue
        if len(tok) >= 2 and tok.startswith('`') and tok.endswith('`'):
            run = p.add_run(tok[1:-1])
            set_font(run, 'Consolas', '宋体', size - 0.5, color=RGBColor(0xA0, 0x20, 0x20))
        elif len(tok) >= 4 and tok.startswith('**') and tok.endswith('**'):
            run = p.add_run(tok[2:-2])
            set_font(run, 'Times New Roman', '黑体', size, bold=True, color=base_color)
        else:
            run = p.add_run(tok)
            set_font(run, 'Times New Roman', '宋体', size, color=base_color)


# ----------------------------------------------------------------------------
# block helpers
# ----------------------------------------------------------------------------

def heading(doc, text, level):
    h = doc.add_heading('', level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = h.add_run(text)
    sizes = {0: 22, 1: 15, 2: 12.5, 3: 11}
    set_font(run, 'Times New Roman', '黑体', sizes.get(level, 11), bold=True, color=NAVY)
    h.paragraph_format.space_before = Pt(10 if level <= 1 else 6)
    h.paragraph_format.space_after = Pt(4)
    return h


def para(doc, text, size=10.5, align=None, space_after=5):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.28
    add_rich(p, text, size)
    return p


def li(doc, text, size=10.5, indent=0.8):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.22
    add_rich(p, text, size)
    return p


def code_block(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Cm(0.4)
    pf.space_before = Pt(3)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.05
    shade_par(p, CODE_FILL)
    lines = text.split('\n')
    for i, line in enumerate(lines):
        if i > 0:
            p.add_run().add_break()
        run = p.add_run(line if line else ' ')
        set_font(run, 'Consolas', '宋体', 9, color=RGBColor(0x1A, 0x1A, 0x1A))
    return p


def table_block(doc, headers, rows, caption=None, col_widths=None, font=9.0):
    if caption:
        TBL[0] += 1
        cp = doc.add_paragraph()
        cp.paragraph_format.space_before = Pt(6)
        cp.paragraph_format.space_after = Pt(2)
        cr = cp.add_run("表 %d  %s" % (TBL[0], caption))
        set_font(cr, 'Times New Roman', '黑体', 10, bold=True, color=NAVY)
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = CENTER
        run = p.add_run(h)
        set_font(run, 'Times New Roman', '黑体', font + 0.5, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        shade_cell(cell, HEADER_FILL)
    # body
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(1)
            add_rich(p, str(val), font)
            if ri % 2 == 1:
                shade_cell(cell, "F2F5FB")
    if col_widths:
        for i, w in enumerate(col_widths):
            for r in t.rows:
                r.cells[i].width = Cm(w)
    return t


def _img_width_cm(path, default=13.5):
    try:
        with Image.open(path) as im:
            w, h = im.size
        aspect = w / float(h)
    except Exception:
        return default
    if aspect >= 1.15:
        return 13.5
    if aspect >= 0.85:
        return 10.0
    return 7.2


def fig(doc, path, caption, width_cm=None):
    if not os.path.exists(path):
        print("WARN missing image:", path)
        para(doc, "[缺图] %s" % os.path.basename(path), size=9)
        return
    if width_cm is None:
        width_cm = _img_width_cm(path)
    p = doc.add_paragraph()
    p.alignment = CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run()
    run.add_picture(prep(path), width=Cm(width_cm))
    FIG[0] += 1
    cap = doc.add_paragraph()
    cap.alignment = CENTER
    cap.paragraph_format.space_after = Pt(8)
    cr = cap.add_run("图 %d  %s" % (FIG[0], caption))
    set_font(cr, 'Times New Roman', '宋体', 9, color=GRAY)


def fig_grid(doc, items, cols=2, width_cm=7.0):
    """items: list of (path, caption). Lays out a borderless grid with captions."""
    rows = (len(items) + cols - 1) // cols
    t = doc.add_table(rows=rows, cols=cols)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    idx = 0
    for r in range(rows):
        for c in range(cols):
            cell = t.cell(r, c)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            if idx < len(items):
                path, capt = items[idx]
                p = cell.paragraphs[0]
                p.alignment = CENTER
                p.paragraph_format.space_after = Pt(1)
                if os.path.exists(path):
                    run = p.add_run()
                    run.add_picture(prep(path), width=Cm(width_cm))
                else:
                    print("WARN missing image:", path)
                    add_rich(p, "[缺图] %s" % os.path.basename(path), 8)
                FIG[0] += 1
                cp = cell.add_paragraph()
                cp.alignment = CENTER
                cp.paragraph_format.space_after = Pt(6)
                cr = cp.add_run("图 %d  %s" % (FIG[0], capt))
                set_font(cr, 'Times New Roman', '宋体', 8.5, color=GRAY)
            idx += 1
    return t


def callout(doc, title, lines):
    """A shaded single-cell note box."""
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    shade_cell(cell, CALL_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_font(r, 'Times New Roman', '黑体', 10.5, bold=True, color=NAVY)
    for ln in lines:
        pp = cell.add_paragraph()
        pp.paragraph_format.space_after = Pt(1)
        pp.paragraph_format.line_spacing = 1.2
        add_rich(pp, ln, 10)
    # spacing after box
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# ----------------------------------------------------------------------------
# cover / TOC / header helpers
# ----------------------------------------------------------------------------

def spacer(doc, pts):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(" ")
    set_font(r, 'Times New Roman', '宋体', pts)
    return p


def center_image(doc, path, width_cm):
    p = doc.add_paragraph()
    p.alignment = CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if os.path.exists(path):
        p.add_run().add_picture(prep(path), width=Cm(width_cm))
    else:
        print("WARN missing asset:", path)
    return p


def cell_bottom_border(cell, sz=6, color='404040'):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(sz))
    bottom.set(qn('w:space'), '0')
    bottom.set(qn('w:color'), color)
    borders.append(bottom)
    tcPr.append(borders)


def para_bottom_border(p, sz=6, color='B0B0B0'):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(sz))
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color)
    pbdr.append(bottom)
    pPr.append(pbdr)


def add_field(paragraph, instr, default_text=''):
    run = paragraph.add_run()
    b = OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'), 'begin')
    i = OxmlElement('w:instrText'); i.set(qn('xml:space'), 'preserve'); i.text = instr
    s = OxmlElement('w:fldChar'); s.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve'); t.text = default_text
    e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'end')
    for el in (b, i, s, t, e):
        run._element.append(el)
    return run


def add_toc(doc):
    p = doc.add_paragraph()
    add_field(p, ' TOC \\o "1-3" \\h \\z \\u ',
              '（在 Word 中按 Ctrl+A 再按 F9 更新此目录域）')


def set_pgnum_start(section, start=1):
    sectPr = section._sectPr
    old = sectPr.find(qn('w:pgNumType'))
    if old is not None:
        sectPr.remove(old)
    pg = OxmlElement('w:pgNumType')
    pg.set(qn('w:start'), str(start))
    sectPr.append(pg)
doc = Document()
normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(10.5)
normal._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

sec = doc.sections[0]
sec.page_height = Mm(297)
sec.page_width = Mm(210)
sec.top_margin = Cm(2.3)
sec.bottom_margin = Cm(2.0)
sec.left_margin = Cm(2.4)
sec.right_margin = Cm(2.4)

# ============================================================================
# COVER PAGE  (section 0 — no header / footer / page number)
# ============================================================================
spacer(doc, 14)
center_image(doc, GDUT_LOGO, 13.8)
spacer(doc, 34)

ct = doc.add_paragraph(); ct.alignment = CENTER; ct.paragraph_format.space_after = Pt(8)
r = ct.add_run("FPGA 课程设计报告")
set_font(r, 'Times New Roman', '黑体', 30, bold=True, color=RGBColor(0, 0, 0))

cs = doc.add_paragraph(); cs.alignment = CENTER; cs.paragraph_format.space_after = Pt(6)
r = cs.add_run("《基于 ZYNQ7020 的图像处理系统设计》")
set_font(r, 'Times New Roman', '黑体', 16, bold=True, color=RGBColor(0, 0, 0))

spacer(doc, 26)

# underlined fill-in fields (姓名/学号 留作占位，用户自行填写)
fields = [
    ("姓　　名", "宋坤峰"),
    ("学　　号", "3123009575"),
    ("专　　业", "微电子科学与工程"),
    ("班　　级", "23 级一班"),
    ("队　　名", "组名是必须"),
    ("指导老师", "周贤中"),
]
ft = doc.add_table(rows=len(fields), cols=2)
ft.alignment = WD_TABLE_ALIGNMENT.CENTER
for ri, (k, v) in enumerate(fields):
    lc = ft.rows[ri].cells[0]
    vc = ft.rows[ri].cells[1]
    lc.width = Cm(3.4)
    vc.width = Cm(6.4)
    lp = lc.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    lp.paragraph_format.space_after = Pt(12)
    rr = lp.add_run(k + "：")
    set_font(rr, 'Times New Roman', '黑体', 14, bold=True)
    vp = vc.paragraphs[0]
    vp.alignment = CENTER
    vp.paragraph_format.space_after = Pt(12)
    rv = vp.add_run(v)
    set_font(rv, 'Times New Roman', '楷体', 14)
    cell_bottom_border(vc, sz=6, color='404040')

spacer(doc, 30)
center_image(doc, IC_LOGO, 5.4)
spacer(doc, 8)
cd = doc.add_paragraph(); cd.alignment = CENTER
r = cd.add_run("2026 年 6 月")
set_font(r, 'Times New Roman', '楷体', 14)

# ============================================================================
# SECTION BREAK -> body section (running header + page numbers from 1)
# ============================================================================
body_sec = doc.add_section(WD_SECTION.NEW_PAGE)
body_sec.page_height = Mm(297)
body_sec.page_width = Mm(210)
body_sec.top_margin = Cm(2.3)
body_sec.bottom_margin = Cm(2.0)
body_sec.left_margin = Cm(2.4)
body_sec.right_margin = Cm(2.4)

# running header: 广东工业大学 │ FPGA 课程设计报告  ........  <当前章节>
body_sec.header.is_linked_to_previous = False
hp = body_sec.header.paragraphs[0]
hp.paragraph_format.tab_stops.add_tab_stop(Cm(16.2), WD_TAB_ALIGNMENT.RIGHT)
rh1 = hp.add_run("广东工业大学")
set_font(rh1, 'Times New Roman', '黑体', 9, bold=True, color=RGBColor(0xB0, 0x12, 0x16))
rh2 = hp.add_run("  │  FPGA 课程设计报告")
set_font(rh2, 'Times New Roman', '宋体', 9, color=GRAY)
hp.add_run("\t")
rh3 = hp.add_run("基于 ZYNQ7020 的图像处理系统设计")
set_font(rh3, 'Times New Roman', '宋体', 9, color=GRAY)
para_bottom_border(hp, sz=4, color='C0C0C0')

# footer: centered page number
body_sec.footer.is_linked_to_previous = False
fpp = body_sec.footer.paragraphs[0]
fpp.alignment = CENTER
pr = add_field(fpp, ' PAGE \\* MERGEFORMAT ', '1')
set_font(pr, 'Times New Roman', '宋体', 9, color=GRAY)

# cover section stays header/footer free
doc.sections[0].header.is_linked_to_previous = False
doc.sections[0].footer.is_linked_to_previous = False
set_pgnum_start(body_sec, 1)

# ============================================================================
# 摘要 + 关键词
# ============================================================================
ah = doc.add_heading('', level=1)
ah.alignment = CENTER
ah.paragraph_format.space_before = Pt(4)
ah.paragraph_format.space_after = Pt(8)
r = ah.add_run("摘　要")
set_font(r, 'Times New Roman', '黑体', 15, bold=True, color=NAVY)

ap = doc.add_paragraph()
ap.paragraph_format.line_spacing = 1.45
ap.paragraph_format.first_line_indent = Cm(0.74)
ap.paragraph_format.space_after = Pt(6)
add_rich(ap,
         "本文面向黑金 ZYNQ7020 开发板，完成一个由浅入深的 FPGA 图像处理系统：从 RTL 级 Sobel 仿真，"
         "到 HDMI 固定图显示、PC 串口传图、PS 经 AXI BRAM 与 PL 共享图像，最终实现由上位机命令控制的"
         "多模式 HDMI 显示（原图 / 灰度 / Sobel 边缘 / 彩色叠加，阈值可调）。在此基础上完成两个相互正交的"
         "第二周综合扩展：综合扩展任务 1（C 档）在上位机端把任意输入尺寸统一适配为固定 `128×72 RGB888`，"
         "覆盖 `128×72`、`160×90`、`144×108` 三种参考尺寸而不改动任何硬件；综合扩展任务 3（B 档）在 PL 上"
         "新增图像锐化算法，复用已验证的 Sobel 3×3 窗口附加 4 邻域拉普拉斯，强度经控制字实时可调，"
         "并以与硬件逐位一致的软件 golden 通过无板软硬件协同仿真逐像素验证。系统经 Vivado 2023.2 综合实现，"
         "并已现场上板验收，HDMI 显示稳定。", 10.5)

kw = doc.add_paragraph()
kw.paragraph_format.space_after = Pt(4)
r = kw.add_run("关键词：")
set_font(r, 'Times New Roman', '黑体', 10.5, bold=True)
r = kw.add_run("ZYNQ7020；FPGA；Sobel 边缘检测；图像锐化；HDMI 显示；上位机；软硬件协同仿真")
set_font(r, 'Times New Roman', '宋体', 10.5)

# ============================================================================
# 目录 (TOC field — Word/导出 PDF 时自动生成)
# ============================================================================
spacer(doc, 8)
th = doc.add_paragraph()
th.alignment = CENTER
th.paragraph_format.space_after = Pt(8)
r = th.add_run("目　录")
set_font(r, 'Times New Roman', '黑体', 18, bold=True, color=NAVY)
add_toc(doc)
pbk = doc.add_paragraph()
pbk.add_run().add_break(WD_BREAK.PAGE)

# ============================================================================
# 1. 课程设计任务说明
# ============================================================================
heading(doc, "1. 课程设计任务说明", 1)
para(doc, "本课程设计以黑金 ZYNQ7020 开发板为平台，完成一个由浅入深的 FPGA 图像处理系统。系统从 RTL 级 Sobel 仿真开始，逐步完成 HDMI 固定图显示、固定图 Sobel、PC 串口传图、PS 写入 AXI BRAM、PL 读取 BRAM 并 HDMI 显示、UART 输入图像 Sobel 显示，以及上位机命令控制显示模式。")

para(doc, "按照老师仓库最新的报告要求，最终报告聚焦第二周工作（实验 5 上位机控制显示 + 综合扩展），第一周内容已在仿真报告和初步实验报告中提交；为保证报告自洽，本文仍对基础实验与第一周扩展做简要回顾。基础任务包括：")
for s in [
    "1. 完成 `sobel_00_rtl_sim` RTL 仿真，说明输入图像、输出边缘图和关键波形。",
    "2. 完成 `sobel_01_hdmi_pattern` HDMI 固定图片显示，验证 HDMI 输出链路。",
    "3. 完成 `sobel_02_hdmi_sobel` 固定图片 Sobel 显示，说明灰度转换和 Sobel 数据流。",
    "4. 完成 `sobel_03_uart_hdmi`，通过串口接收 PC 图像并写入 BRAM。",
    "5. 使用 `host_camera_uart` 的 GUI 或命令行脚本发送图像到 ZYNQ。",
    "6. 完成 `sobel_04_uart_sobel_hdmi`，在 HDMI 上显示 UART 输入图像的 Sobel 边缘结果。",
    "7. 完成 `sobel_05_pc_control_display`，通过 PC 控制显示模式、Sobel 阈值和边缘叠加。",
    "8. 实验 0 到实验 4 每个实验至少完成 1 个基础扩展；实验 5 完成上位机控制显示能力。",
    "9. 保存上板现象照片、串口输出、Vivado 资源利用率、时序和 DRC 结果。",
]:
    li(doc, s)

para(doc, "第二周综合扩展老师给出“三选一”。本组在完成 **综合扩展任务 1（C 档，纯 PC 端输入规格扩展）** 的基础上，进一步完成 **综合扩展任务 3（B 档，增加 PL 图像处理算法——图像锐化）**，分别整理到独立目录 `大拓展_01_上位机输入规格扩展` 与 `大拓展_03_图像处理算法扩展`，不与实验 5 基础工程混合。两个扩展互相正交、一软一硬：任务 1 在输入侧做 PC 端规格适配（不改硬件），任务 3 在算法侧给 PL 新增锐化（改 PL/PS 但原 Sobel 零回归）；二者共用同一上位机工具与同一 bitstream，可叠加演示（先缩放、后锐化），属超额完成。任务 1 详见第 6 节，任务 3 详见第 7 节。")

heading(doc, "1.1 完成度与老师等级评定对照", 2)
para(doc, "下表把本组成果直接对照老师仓库最新加入的“等级评定要求”。两个扩展的目标档位均已达成，并有现场上板实拍与逐像素软件对比为证。")
table_block(
    doc,
    ["综合扩展", "老师等级要求（目标档）", "本组达成情况", "现场证据"],
    [
        ["任务 1 输入规格扩展", "C 档：硬件不变，上位机统一缩放至 ≥3 种尺寸（128×72、160×90、144×108）后发送", "已达成 C 档：三种参考尺寸全部实现并自动化矩阵验证，硬件复用实验 5", "图 128×72 / 160×90 / 144×108 上板实拍"],
        ["任务 3 图像处理算法", "B 档：新增 1 种算法 + 上板演示 + 与软件参考对比", "已达成 B 档：新增图像锐化，上板实时调强度，11 配置逐像素 =match", "上位机软件预览 + 上板锐化实拍"],
    ],
    caption="综合扩展完成度与老师等级评定对照",
    col_widths=[3.2, 5.6, 4.8, 2.8],
)

# ============================================================================
# 2. 系统总体结构和数据流
# ============================================================================
heading(doc, "2. 系统总体结构和数据流", 1)
para(doc, "系统整体分为 PC 上位机、ZYNQ PS、AXI BRAM、PL 图像处理和 HDMI 输出五个部分，主链路如下：")
code_block(doc,
           "PC 上位机(图片/摄像头/视频/控制命令)\n"
           "   -> UART 图像帧 / 控制帧\n"
           "      -> ZYNQ PS(协议解析)\n"
           "         -> AXI BRAM(图像区 + 控制字)\n"
           "            -> PL(灰度 / Sobel / 锐化 / 显示选择)\n"
           "               -> HDMI 显示(原图 / 灰度 / 边缘 / 叠加 / 锐化)")
para(doc, "两个综合扩展的接入点：任务 1 在“PC 上位机 -> UART 图像帧”之间做尺寸适配，硬件侧契约保持固定 `128×72 RGB888`；任务 3 在“PL 图像处理”内新增锐化支路，并经控制字实时调强度。基础数据流如下表。")
table_block(
    doc,
    ["实验", "输入来源", "主要处理路径", "输出现象"],
    [
        ["实验 0", "input_rgb.hex", "UART 模型 → 帧解析 → rgb_to_gray → sobel_core", "PNG/PGM 边缘图与关键波形"],
        ["实验 1", "ROM 固定图", "ROM → HDMI 时序 → rgb2dvi", "HDMI 固定图放大显示"],
        ["实验 2", "ROM 固定图", "ROM → 灰度 → Sobel → 阈值二值化", "HDMI 黑白边缘图"],
        ["实验 3", "PC UART 图像", "PC → UART → PS → BRAM → PL 读图", "HDMI 原图显示"],
        ["实验 4", "PC UART 图像", "PC → UART → PS → BRAM → PL Sobel", "HDMI 彩色边缘图"],
        ["实验 5", "PC 图像 + 控制帧", "PC 控制 → PS 控制字 → PL 显示模式选择", "原图/灰度/边缘/叠加切换"],
        ["综合扩展", "任意尺寸图片/目录/视频", "PC 端缩放裁剪补边 → 128×72 等处理尺寸 → 固定 128×72 帧；PL 新增锐化", "覆盖 C 档 3 种尺寸 + 实时锐化，不改硬件契约"],
    ],
    caption="各实验输入来源、处理路径与输出现象",
    col_widths=[1.8, 3.0, 6.6, 4.2],
)
para(doc, "UART 图像帧保持 RGB888 协议：")
code_block(doc,
           "frame header: 55 aa width_l width_h height_l height_h 18\n"
           "line header : 33 cc row_l row_h\n"
           "pixels      : R G B, repeated 128 times per row")
para(doc, "实验 5 控制帧格式为：")
code_block(doc,
           "A5 5A cmd value\n"
           "cmd=1: mode, cmd=2: threshold, cmd=3: overlay, cmd=4: sharpen(任务3新增)")

# ============================================================================
# 3. 基础实验复现过程
# ============================================================================
heading(doc, "3. 基础实验复现过程（回顾）", 1)
para(doc, "本节为基础链路回顾，详细仿真与第一周扩展已在仿真报告、初步实验报告中提交。下列现场照片用于说明各级硬件链路均已稳定上板。")

heading(doc, "3.1 实验 0：RTL 仿真", 2)
para(doc, "实验 0 完成纯 RTL 仿真，验证 UART 帧解析、RGB 转灰度、Sobel 边缘检测与输出帧完成信号。ModelSim 输出 `Sobel RGB888 simulation passed`，有效像素 `128×72=9216`，Sobel 非零像素 6980 个，最大值 255。")
fig_grid(doc, [
    (EV + "/01_rtl_sim/exp00_input_rgb.png", "实验 0 输入图"),
    (EV + "/01_rtl_sim/exp00_sobel_out.png", "实验 0 Sobel 输出"),
], cols=2, width_cm=6.0)
fig(doc, EV + "/01_rtl_sim/exp00_key_waveform.png", "实验 0 关键波形（帧解析、灰度、Sobel、完成信号）")

heading(doc, "3.2 实验 1：HDMI 固定图显示", 2)
para(doc, "实验 1 验证 HDMI 输出链路，固定 `128×72` 图像放大为 `1280×720` 显示，蓝色边框标记有效显示区域。")
fig(doc, EV + "/02_hdmi_pattern/exp01_hdmi_field_20260618.jpg", "实验 1 HDMI 固定图现场（蓝色有效区域边框）", width_cm=11.0)

heading(doc, "3.3 实验 2：固定图 Sobel HDMI 显示", 2)
para(doc, "实验 2 在固定图基础上加入灰度转换和 Sobel 运算，默认阈值 80，HDMI 显示黑白边缘图。阈值 40/80/120 的白色边缘像素分别为 1307/1274/1234，随阈值升高单调减少。")
fig(doc, EV + "/03_hdmi_sobel/exp02_hdmi_sobel_field_20260618.jpg", "实验 2 固定图 Sobel 现场", width_cm=11.0)

heading(doc, "3.4 实验 3：UART 传图并 HDMI 显示原图", 2)
para(doc, "实验 3 用 PC 工具发送图像，PS 解析 UART 帧写入 BRAM，PL 读取后 HDMI 显示原图，验证 PC/UART/PS/AXI BRAM/PL HDMI 的基础闭环。")
fig(doc, EV + "/04_uart_hdmi/exp03_uart_hdmi_field_20260618.jpg", "实验 3 UART 原图显示现场", width_cm=11.0)

heading(doc, "3.5 实验 4：UART 输入图像 Sobel 显示", 2)
para(doc, "实验 4 在实验 3 基础上加入 PL Sobel，默认绿色边缘、黑色背景，支持阈值对比与不同输入图像测试。")
fig_grid(doc, [
    (EV + "/05_uart_sobel/exp04_uart_sobel_field_20260618.jpg", "实验 4 UART Sobel 显示"),
    (EV + "/05_uart_sobel/exp04_uart_sobel_natural_image_field_20260618.jpg", "实验 4 自然图像边缘"),
    (EV + "/05_uart_sobel/exp04_uart_sobel_strong_edges_field_20260618.jpg", "实验 4 强边缘图像"),
], cols=3, width_cm=4.7)

heading(doc, "3.6 实验 5：上位机控制显示模式", 2)
para(doc, "实验 5 增加 PC 控制帧，支持原图、灰度、Sobel 边缘、红色边缘叠加等显示模式，并支持阈值调整。这是第二周综合扩展的基础链路，老师明确要求最终报告展示其模式切换、阈值控制与彩色叠加效果。")
fig_grid(doc, [
    (EV + "/06_pc_control/exp05_mode_original_field_20260618.jpg", "实验 5 原图模式"),
    (EV + "/06_pc_control/exp05_mode_gray_field_20260618.jpg", "实验 5 灰度模式"),
    (EV + "/06_pc_control/exp05_mode_edge_threshold_sparse_field_20260618.jpg", "实验 5 边缘模式"),
    (EV + "/06_pc_control/exp05_mode_overlay_field_20260618.jpg", "实验 5 红色叠加模式"),
], cols=2, width_cm=7.0)
para(doc, "实验 5 还支持 Sobel 阈值的实时调节。下面三张为同一边缘画面在不同阈值下的 HDMI 显示效果：阈值越高，检出的白色边缘越稀疏、越干净；阈值越低，白色边缘越密集、保留的细节越多——可按场景在“多检出”与“少噪声”之间权衡。这直接对应老师对实验 5“阈值控制”的展示要求。")
fig_grid(doc, [
    (ASSETS + "/thr_sparse.png", "阈值较高：白色边缘最稀疏"),
    (ASSETS + "/thr_mid.png", "阈值中等：边缘密度居中"),
    (ASSETS + "/thr_dense.png", "阈值较低：白色边缘最密集"),
], cols=3, width_cm=4.9)

# ============================================================================
# 4. 第一周基础扩展完成情况
# ============================================================================
heading(doc, "4. 第一周基础扩展完成情况", 1)
table_block(
    doc,
    ["实验", "基础扩展", "完成现象"],
    [
        ["实验 0", "异常帧自检", "testbench 覆盖错误帧头、错误格式、错误行号与合法 RGB888 帧"],
        ["实验 1", "HDMI 有效区域边框", "固定图周围增加蓝色边框，确认显示区域与坐标映射"],
        ["实验 2", "Sobel 阈值对比", "阈值 40/80/120 下边缘像素单调减少"],
        ["实验 3", "原图显示边框", "hdmi_bram_display.v 增加 BORDER_WIDTH/BORDER_COLOR，不修改 BRAM 数据"],
        ["实验 4", "彩色边缘与阈值对比", "默认绿色边缘，阈值 40/80/120 下边缘数量单调减少"],
        ["实验 5", "上位机控制显示", "PC 控制原图、灰度、边缘、叠加、阈值与 overlay"],
    ],
    caption="第一周基础扩展完成情况",
    col_widths=[1.8, 3.6, 10.2],
)
fig(doc, EV + "/02_hdmi_pattern/exp01_hdmi_gradient_color_bars_field_20260618.jpg", "实验 1 渐变/色条现场（说明 HDMI 输出稳定性）", width_cm=10.5)

# ============================================================================
# 5. 修改文件和关键代码说明
# ============================================================================
heading(doc, "5. 修改文件和关键代码说明", 1)
heading(doc, "5.1 基础实验关键模块", 2)
table_block(
    doc,
    ["模块或文件", "所在实验", "功能"],
    [
        ["rgb_to_gray.v", "0、2、4、5", "使用 (77R+150G+29B)>>8 完成 RGB888 到灰度转换"],
        ["sobel_core.v", "0、2、4、5", "3×3 窗口 + 两行缓存计算 Sobel 梯度，输出 8 bit 边缘强度"],
        ["image_frame_rx.v", "0", "RTL 仿真中解析 UART 图像帧和行头"],
        ["hdmi_image_display.v", "1、2", "生成 HDMI 有效区域图像并完成 10× 放大显示"],
        ["hdmi_bram_display.v", "3、5", "从 BRAM 读取 RGB888 原图并输出 HDMI"],
        ["hdmi_bram_sobel_display.v", "4、5", "BRAM 读图后完成灰度、Sobel、阈值、颜色映射与显示模式选择"],
        ["main.c", "3、4、5", "PS 端解析 UART 图像帧/控制帧，写入 BRAM 图像区和控制字"],
        ["camera_uart_sender.py / camera_uart_gui.py", "PC 工具", "读取图片/摄像头/视频并发送 RGB888 图像帧和控制帧"],
    ],
    caption="基础实验关键模块",
    col_widths=[5.6, 2.2, 7.8],
)

heading(doc, "5.2 综合扩展一（任务 1 输入规格）修改文件", 2)
para(doc, "任务 1 已从 `exp/05-ext-scaling` 整理为独立目录 `大拓展_01_上位机输入规格扩展`，只改 PC 上位机；实验 5 工程作为基础工程保留。")
table_block(
    doc,
    ["文件", "修改内容"],
    [
        ["大拓展_01.../host_tool/camera_uart_sender.py", "新增 prepare_frame、--fit-mode、--proc-size、--fill-color；支持 stretch/letterbox/center-crop；输出帧固定 128×72"],
        ["大拓展_01.../host_tool/camera_uart_gui.py", "GUI 增加 Fit mode、Proc size 下拉框与 Folder 目录入口；GUI 与 CLI 共用 prepare_frame"],
        ["大拓展_01.../README.md", "说明 C/B/A 评分边界、运行方式与独立交付目录结构"],
        ["大拓展_01.../host_tool/tests/test_prepare_frame.py", "测试缩放/裁剪/补边策略、输出形状、C 档参考尺寸覆盖与兼容旧 resize"],
        ["大拓展_01.../host_tool/tests/test_protocol_invariants.py", "测试图像帧包长、帧头与控制帧字节不变"],
        ["大拓展_01.../evidence/*", "保存离线验证截图、测试日志、CLI help 与矩阵结果"],
    ],
    caption="综合扩展一修改文件",
    col_widths=[6.6, 9.0],
)
para(doc, "关键实现思想：")
code_block(doc,
           "任意尺寸 BGR 输入\n"
           "  -> prepare_frame(frame_bgr, width=128, height=72, fit_mode, content_size, fill)\n"
           "  -> 固定 128x72 RGB888\n"
           "  -> build_frame_packet\n"
           "  -> UART 发送给原 sobel_05 链路")

heading(doc, "5.3 综合扩展二（任务 3 锐化）修改文件", 2)
para(doc, "任务 3 是“改 PL 加算法”，涉及 PL/PS/上位机三侧，但改动是 **纯附加**：原 Sobel 四模式（原图/灰度/边缘/叠加）逐位不变、零回归。独立交付目录 `大拓展_03_图像处理算法扩展`。")
table_block(
    doc,
    ["文件", "层", "修改内容"],
    [
        ["sobel_core.v", "PL", "在已验证 3×3 窗口上附加输出 4 邻域拉普拉斯 lap_data（中心=mid1），Sobel 端口与时序不变"],
        ["hdmi_bram_sobel_display.v", "PL", "新增 MODE_SHARPEN(4)、lap_mem、控制字 0x900C；显示端 clamp(rgb+(k·lap)>>>8)；display_mode 位宽 2→3"],
        ["main.c", "PS", "新增控制命令 0x04（锐化强度→0x900C）；mode 掩码 0x03→0x07"],
        ["大拓展_03.../host_tool/sharpen_algo.py", "PC", "锐化软件 golden（numpy），与 RTL 定点逐位一致"],
        ["大拓展_03.../host_tool/camera_uart_sender.py", "PC", "锐化感知发送器（sharpen 模式 + 0x04 命令），并入任务 1 的 prepare_frame/fit/proc 缩放"],
        ["大拓展_03.../host_tool/camera_uart_gui.py", "PC", "验收 GUI：模式选择 + 实时锐化强度滑块 + 原图/软件锐化并排预览 + 任务 1 的 fit/proc 控件"],
        ["tools/generate_exp05_expected.py、tools/cosim/*", "仿真", "golden 库新增 laplacian()+MODE_SHARPEN；协同仿真新增 4 档锐化逐像素对比与真实 PS 命令端到端校验"],
    ],
    caption="综合扩展二修改文件",
    col_widths=[5.2, 1.3, 9.1],
)

# ============================================================================
# 6. 综合扩展一
# ============================================================================
heading(doc, "6. 第二周综合扩展（一）：上位机与输入规格扩展（任务 1，C 档）", 1)
heading(doc, "6.1 扩展题目和设计方案", 2)
para(doc, "选择综合扩展任务 1：基于 `sobel_05` 的上位机与输入规格扩展。采用方案 A：上位机统一完成尺寸适配，FPGA 接收侧保持固定 `128×72 RGB888`。好处是：")
for s in [
    "1. 不修改 BRAM 地址映射，不改变 PS 接收协议。",
    "2. 不修改 PL Sobel 处理尺寸，不影响 HDMI 放大逻辑。",
    "3. 兼容实验 5 现有显示控制命令（原图、灰度、边缘、叠加、阈值）。",
    "4. 可快速验证多输入尺寸、多输入来源与不同缩放策略。",
]:
    li(doc, s)
callout(doc, "对照老师等级评定（任务 1）→ 达成 C 档", [
    "C 档要求：硬件不变，上位机统一缩放至 ≥3 种尺寸（参考 128×72、160×90、144×108）后发送。",
    "本组实现：三种参考尺寸全部支持并进入自动化矩阵验证，硬件复用实验 5 bitstream，零硬件改动。上板实拍见 6.3 节图。",
])

heading(doc, "6.2 离线验证结果", 2)
para(doc, "离线测试覆盖 `640×480`、`1920×1080`、`1080×1920` 三种原始尺寸，三种 fit 模式与四种处理尺寸，形成 `3×3×4=36` 组验证。四种处理尺寸为 `128x72`、`160x90`、`144x108`、`64x36`，其中前三者即评分表 C 档参考尺寸。所有输出最终均为 `(72,128,3)` uint8，打包后帧长均为 `27943` 字节，帧头均为 `55 aa 80 00 48 00 18`。")
para(doc, "三种 fit 策略对比：`stretch` 直接拉伸铺满（可能变形）；`letterbox` 等比缩放居中补边（不变形）；`center-crop` 等比放大居中裁剪（无黑边）。")
fig_grid(doc, [
    (NEW + "/strech_128x72.png", "stretch 128×72（拉伸铺满）"),
    (NEW + "/letterbos_128x72.png", "letterbox 128×72（补边不变形）"),
    (NEW + "/centercrop_128x72.png", "center-crop 128×72（裁剪铺满）"),
], cols=3, width_cm=4.7)
para(doc, "处理尺寸对比中，`64×36` 作为更低处理尺寸会产生更明显的块状效果，但最终仍放回固定 `128×72` 帧，UART 帧长和帧头不变。")
fig_grid(doc, [
    (NEW + "/strech_64x36.png", "stretch 64×36（块状更明显）"),
    (NEW + "/centrecrop_64x36.png", "center-crop 64×36"),
], cols=2, width_cm=6.0)
para(doc, "离线测试结果：")
code_block(doc,
           "tests/test_prepare_frame.py: 8/8 passed\n"
           "tests/test_protocol_invariants.py: 4/4 passed\n"
           "summary: 12/12 offline assertions passed")
for s in [
    "• `stretch` 与旧工具的 `cv2.resize + BGR->RGB` 结果兼容。",
    "• `letterbox` 保持原图宽高比并用指定颜色填充空白。",
    "• `center-crop` 铺满输出画面，不引入黑边。",
    "• C 档尺寸验收依据：`128x72`、`160x90`、`144x108` 全部实现并通过矩阵测试；`64×36` 仅作额外对比。",
    "• 控制帧仍为 `A5 5A cmd value`，不影响实验 5 显示控制功能。",
]:
    li(doc, s)

heading(doc, "6.3 上板结果与最终验收截图", 2)
para(doc, "硬件侧复用实验 5 的 bitstream 与 PS 程序。由于扩展不修改 FPGA 处理尺寸，现场上板重点验证三种 C 档参考处理尺寸与原 `sobel_05` 显示链路是否仍能稳定工作。下列为 2026-06-26 现场验收实拍：同一画面分别以 `128×72`、`160×90`、`144×108` 处理尺寸送入上位机，最终统一映射回固定 `128×72` 帧后由 HDMI 显示，可见不同处理尺寸带来的块状/细节差异，而硬件链路完全不变。")
fig(doc, NEW + "/128x72.png", "任务 1 上板实拍：处理尺寸 128×72（C 档参考尺寸，HDMI 实拍）", width_cm=12.5)
fig(doc, NEW + "/160x90.png", "任务 1 上板实拍：处理尺寸 160×90（C 档参考尺寸）", width_cm=12.5)
fig(doc, NEW + "/144x108.png", "任务 1 上板实拍：处理尺寸 144×108（C 档参考尺寸）", width_cm=12.5)
para(doc, "结论：第二周扩展（一）能够在上位机端完成不同输入规格到固定 `128×72 RGB888` 的映射，覆盖 C 档要求的三种参考尺寸，并保持实验 5 的 FPGA 硬件链路不变，达成 C 档评定。")

# ============================================================================
# 6B. 综合扩展二
# ============================================================================
heading(doc, "7. 第二周综合扩展（二）：增加 PL 图像处理算法——锐化（任务 3，B 档）", 1)
heading(doc, "7.1 题目与设计方案", 2)
para(doc, "选择综合扩展任务 3：在 `sobel_05_pc_control_display` 的 PL 上新增一种图像处理算法，并通过上位机命令选择显示结果。本组实现 **图像锐化（Laplacian / unsharp 增强）**。设计取向是“最低风险地改 PL 加算法”：")
for s in [
    "1. **复用已验证窗口**：锐化所需的 4 邻域拉普拉斯与 Sobel 共用 `sobel_core.v` 同一条已验证 3×3 滑窗（中心像素=mid1），仅附加输出端口 `lap_data`，Sobel 通路逐位不变，零回归。",
    "2. **不动 BRAM / 时序 / BD**：锐化只新增一块 `lap_mem` 与显示端一次乘加，不改 Block Design、BRAM 容量和 HDMI 时序。",
    "3. **实时可调**：锐化强度 `k` 经控制帧 `A5 5A 04 k` 写入控制字 `0x900C`，PL 每帧读取并在显示端应用，拖动上位机滑块即可让 HDMI 实时变锐，无需重发图像帧——与实验 5 的 mode/threshold/overlay 控制完全同构。",
]:
    li(doc, s)
callout(doc, "对照老师等级评定（任务 3）→ 达成 B 档", [
    "B 档要求：新增 1 种算法 + 上板演示 + 与软件参考对比。",
    "本组实现：新增图像锐化算法并上板（7.5 实拍），强度实时可调；无板协同仿真 11 配置逐像素 `=match`，numpy 软件 golden == RTL 输出（7.3）。",
])

heading(doc, "7.2 锐化定点算法（RTL / PS / 上位机三处逐位一致）", 2)
code_block(doc,
           "gray  = (77*R + 150*G + 29*B) >> 8              # 与 rgb_to_gray.v 一致\n"
           "lap   = 4*center - up - down - left - right      # 4 邻域拉普拉斯, 1 像素边界=0\n"
           "delta = floor(strength * lap / 256)              # RTL 算术右移 >>>8\n"
           "out_c = clamp(c + delta, 0, 255)  for c in R,G,B # 逐通道, 保留彩色")
para(doc, "`strength=0` 时输出与原图逐像素相同；强度越大边缘/纹理越锐，过大时高频处饱和裁剪（overshoot / halo，这是 Laplacian 锐化的固有特性，在真实照片上表现为自然的“变清晰”）。同一套整数运算在 PL（`hdmi_bram_sobel_display.v`）、PS（`main.c`）与上位机（`sharpen_algo.py`）三处实现，保证软件参考与硬件输出逐位一致。")

heading(doc, "7.3 无板软硬件协同仿真验证（RTL == 软件 golden）", 2)
para(doc, "复用实验 5 的无板协同仿真链（真实上位机打包 → 真实 `main.c` PS 分发 → XSim 渲染 HDMI → 与软件 golden 逐像素比对），已扩展覆盖锐化，本地 Vivado 2023.2 XSim 全部通过：")
table_block(
    doc,
    ["验证项", "结果"],
    [
        ["真实 main.c 收 A5 5A 04 96 → 控制字 0x900C", "sharpen=96，图像区 + 4 个控制字逐项一致"],
        ["RTL 自检（原 Sobel 四模式自洽）", "EXP05_SELFCHECK_TB=passed，零回归"],
        ["全分辨率 1280×720 逐像素对比", "11 个配置全部 =match，含 sharp0/64/128/255（sharp0 等于原图）"],
        ["上位机 numpy 预览 vs 软件 golden vs RTL", "三者逐位一致"],
        ["协同仿真链", "EXP05_COSIM_CHAIN=passed"],
    ],
    caption="任务 3 无板协同仿真验证结果",
    col_widths=[8.2, 7.4],
)
para(doc, "这满足任务 3“完成 RTL 仿真，对比硬件输出和软件参考结果”的要求；软件参考即 `sharpen_algo.py`，与 GUI 实时预览同源。")

heading(doc, "7.4 资源利用率与时序（OOC 综合）", 2)
para(doc, "对 PL 显示路径（`hdmi_bram_sobel_display + rgb_to_gray + sobel_core`）做 out-of-context 综合（xc7z020clg400-2，74.25 MHz 像素时钟约束）：")
table_block(
    doc,
    ["指标", "结果"],
    [
        ["时序", "All timing constraints met，WNS +2.452 ns，0 失败端点"],
        ["Slice LUT", "9229（17.35%）"],
        ["Slice Register", "2725（2.56%）"],
        ["Block RAM", "10.5 tile（7.5% / 140）"],
        ["DSP", "4（1.82% / 220）"],
    ],
    caption="任务 3 PL 显示路径 OOC 综合资源与时序",
    col_widths=[5.0, 10.6],
)
para(doc, "锐化新增的 `lap_mem` 仅占数个 BRAM tile，强度乘法 `k·lap` 占 1 个 DSP，组合乘加在 74.25 MHz 下时序裕量充足（约 18%），无需流水线。OOC 报告原文见 `大拓展_03.../evidence/ooc_synth_*.rpt`。")

heading(doc, "7.5 上板演示与验收 GUI（现场实拍）", 2)
para(doc, "验收 GUI（`大拓展_03.../host_tool/camera_uart_gui.py`）面向“手操调参看可见改善”：连接串口 → 载入图片 → 发送 → 显示模式选 `sharpen` → 拖动锐化强度滑块，同窗口「原图 | 软件锐化(k)」并排实时刷新，HDMI 硬件画面同步变锐；切回 `edge` 即为原 Sobel，体现“新增算法 vs 原 Sobel”的对比。该 GUI 同时并入任务 1 的 fit / proc 缩放控件，可叠加演示两个综合扩展。")
fig(doc, NEW + "/上位机.png", "验收上位机工具：同一界面集成任务 1（fit 模式 / 处理分辨率）与任务 3（显示模式 sharpen + 锐化强度 k），下方为「原图 | 软件锐化(k)」并排预览（此处 k=0，等于原图）", width_cm=14.0)
para(doc, "拖动锐化强度滑块时，上位机软件预览（≈ HDMI 硬件输出）实时刷新。下面两图为同一画面在较低强度与强度拉满时的软件预览对比，可见右侧锐化结果边缘/纹理增强、强度拉满时出现自然的 overshoot。")
fig(doc, NEW + "/上位机较低锐化处理预览.png", "上位机软件预览：左原图 | 右软件锐化（较低强度 k≈28）", width_cm=13.5)
fig(doc, NEW + "/上位机锐化拉满预览.png", "上位机软件预览：左原图 | 右软件锐化（强度拉满 k=255，可见锐化 overshoot/halo）", width_cm=13.5)
para(doc, "上板 HDMI 实拍同步变锐，与软件预览一致。下面两图为同一画面在较低锐化与锐化拉满时的 HDMI 现场实拍：")
fig_grid(doc, [
    (NEW + "/较低锐化处理实图.png", "上板 HDMI 实拍：较低锐化强度"),
    (NEW + "/锐化拉满实图.png", "上板 HDMI 实拍：锐化强度拉满，边缘/纹理明显增强"),
], cols=2, width_cm=7.4)
para(doc, "结论：任务 3 在 PL 上新增的图像锐化算法已上板演示，强度实时可调，HDMI 硬件输出与软件 golden 逐像素一致，达成 B 档评定；bitstream 复用实验 5 流程（不改 BD/BRAM/时序）。")

# ============================================================================
# 7. 资源利用率、时序结果和性能分析
# ============================================================================
heading(doc, "8. 资源利用率、时序结果和性能分析", 1)
heading(doc, "8.1 Vivado 资源与时序汇总", 2)
table_block(
    doc,
    ["实验", "LUT", "FF", "BRAM", "DSP", "WNS(ns)", "TNS(ns)", "DRC", "说明"],
    [
        ["实验 1 HDMI 固定图", "232", "179", "12", "0", "+7.772", "0", "0 err,1 warn", "纯 PL HDMI"],
        ["实验 2 HDMI Sobel", "2842", "2353", "14", "0", "+1.578", "0", "0 err,1 warn", "加灰度/Sobel/阈值"],
        ["实验 3 UART HDMI", "1775", "1504", "16", "0", "+8.173", "0", "0 viol", "加 PS/AXI BRAM/原图"],
        ["实验 4 UART Sobel", "4446", "3662", "18", "0", "+13.453", "0", "0 err", "加 BRAM Sobel/绿边"],
        ["实验 5 PC 控制显示", "10795", "4113", "20", "0", "+0.325", "0", "0 viol", "加模式/阈值/叠加"],
        ["任务 1 输入规格扩展", "同实验 5", "同上", "同上", "同上", "同上", "0", "同上", "只改 PC，不改 FPGA 工程"],
        ["任务 3 锐化(PL OOC)", "9229", "2725", "10.5", "4", "+2.452", "0", "0", "仅 PL 显示路径 OOC；bitstream 复用实验 5"],
    ],
    caption="各实验 Vivado 资源利用率与时序汇总",
    col_widths=[3.0, 1.5, 1.5, 1.3, 1.0, 1.8, 1.3, 1.6, 2.6],
    font=8.0,
)
para(doc, "从资源趋势看，实验 1 资源主要用于 HDMI 时序、ROM/BRAM 与 TMDS 输出；实验 2/4/5 随 Sobel、BRAM、显示控制与多模式叠加增加，LUT/FF 明显上升。实验 5 WNS `+0.325 ns` 仍满足时序但裕量最小，因此任务 1 不改 PL、任务 3 仅以纯附加方式扩展，避免继续压缩时序裕量。")

heading(doc, "8.2 性能分析", 2)
para(doc, "图像固定为 `128×72 RGB888` 时，每帧协议层字节数为：")
code_block(doc,
           "frame header 7B + line header 72x4B + pixel 128x72x3B\n"
           "= 7 + 288 + 27648\n"
           "= 27943B")
para(doc, "若使用 UART 8N1 与 `115200` baud，理论传输时间约为：")
code_block(doc, "27943B x 10bit / 115200 ~= 2.43 s/frame")
para(doc, "因此系统瓶颈主要在 UART 传输带宽，而非 PL Sobel/锐化本身。实验 5 与两个综合扩展通过 PC 端预处理保持传输帧固定，避免硬件修改带来的时序风险，但不能突破 UART 带宽限制。若要提升实时性，后续可考虑老师给出的网络传输扩展方向（任务 2），用 UDP/TCP 替代 UART。")

# ============================================================================
# 8. 问题记录与总结
# ============================================================================
heading(doc, "9. 问题记录与总结", 1)
heading(doc, "9.1 问题记录", 2)
table_block(
    doc,
    ["问题", "现象", "解决或处理"],
    [
        ["RTL 仿真工具差异", "本机未检测到 Icarus/VVP", "使用 ModelSim 完成仿真，并用 XSim 入口交叉验证"],
        ["VCD 文件过大", "原始波形约 127MB，不适合提交", "提取关键波形 PNG/SVG，报告引用截图"],
        ["UART 传图速度慢", "115200 baud 下单帧约 2.43 s", "使用低 FPS 或单帧测试；性能分析明确瓶颈"],
        ["实验 5 时序裕量较小", "WNS +0.325 ns", "任务 1 不改 PL；任务 3 纯附加，OOC 后裕量约 18%"],
        ["输入尺寸多样", "不同比例直接 resize 会变形", "任务 1 加入 stretch/letterbox/center-crop 三策略"],
        ["硬件尺寸修改风险高", "改 PL 尺寸需同步 BRAM/PS/HDMI/仿真", "采用 PC 端统一缩放，固定 FPGA 侧 128×72 契约"],
        ["锐化软硬件一致性", "需保证 PL/PS/上位机三处结果相同", "统一定点算法 + 协同仿真 11 配置逐像素 =match"],
    ],
    caption="问题记录",
    col_widths=[3.6, 5.4, 6.6],
)

heading(doc, "9.2 总结", 2)
para(doc, "本课程设计完成了从 RTL 仿真到 HDMI 上板显示、从固定图像到 PC UART 输入图像、从单一 Sobel 输出到上位机控制显示模式的完整流程。基础实验验证了 RGB888 图像在 Verilog、UART、BRAM、PS/PL 与 HDMI 链路中的表示与传递，也验证了灰度转换、Sobel 边缘检测、阈值控制与彩色叠加的硬件实现。")
para(doc, "第一周基础扩展完成了 HDMI 边框、Sobel 阈值对比、彩色边缘映射与显示模式控制。第二周综合扩展（一）选择上位机与输入规格扩展，在不修改硬件的前提下支持不同尺寸输入、三种 fit 策略和四种处理尺寸，其中 `128x72`、`160x90`、`144x108` 覆盖 C 档参考尺寸，并通过 12 项离线测试与 36 组矩阵验证保证 UART 帧格式和控制命令不变，现场上板实拍三种尺寸，达成 C 档。")
para(doc, "在此基础上，本组又完成了综合扩展（二）任务 3（B 档）：在 PL 上新增图像锐化算法，复用 Sobel 的 3×3 窗口附加 4 邻域拉普拉斯，不改 BRAM/时序/BD，强度经控制字 `0x900C` 实时可调；无板协同仿真已逐像素证明 RTL 与软件 golden 一致（11 配置含 4 档锐化全部 `=match`），OOC 综合时序达标（WNS +2.452 ns），并已现场上板演示强度由小到大的锐化效果，达成 B 档。两个综合扩展正交、可在同一上位机工具、同一 bitstream 上叠加演示。")
para(doc, "后续若继续改进，可优先考虑两点：一是将 UART 升级为网络传输（任务 2 方向）以提升帧率；二是从任务 3 的 B 档继续向 A 档演进——再增加 Prewitt/Laplacian 边缘、二值化等算法（`lap_mem` 与灰度已就绪，增量很小），凑齐“≥3 种算法”，并在加算法时同步优化流水线与资源，避免实验 5 已经较小的时序裕量继续收紧。")

# ============================================================================
# 10. 成员分工与贡献
# ============================================================================
heading(doc, "10. 成员分工与贡献", 1)
para(doc, "本课程设计由四人小组协作完成，分工与主要贡献如下表。")
table_block(
    doc,
    ["成员", "分工", "主要贡献"],
    [
        ["宋坤峰", "PL / HDMI / Sobel", "HDMI 显示、Sobel 链路、肖像素材；基础实验复现与答辩汇报"],
        ["陈维瀚", "PS / UART / BRAM", "实验 3/4/5 串口接收、BRAM 写入、控制字协议；第一周基础扩展与照片证据"],
        ["刘照君", "PC 上位机与综合扩展", "大拓展一输入规格（fit/proc、GUI/CLI、离线测试）；大拓展二上位机（锐化软件 golden、GUI 滑块与并排预览）；报告"],
        ["李翠怡", "资料与展示", "照片证据、资源时序整理、报告、PPT 制作与最终演示"],
    ],
    caption="成员分工与贡献",
    col_widths=[2.2, 4.0, 9.4],
)

# ----------------------------------------------------------------------------
doc.save(OUT)
print("SAVED:", OUT)
print("figures:", FIG[0], "tables:", TBL[0])
print("sections:", "(cover + body)")
